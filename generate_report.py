#!/usr/bin/env python3
"""Generate Orbit Team Review Report as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2A, 0x8C, 0x7E)  # Orbit teal
    return h

def add_table_from_data(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ==========================================
# COVER PAGE
# ==========================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ORBIT PLATFORM')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x2A, 0x8C, 0x7E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('First Team Review & Strategic Roadmap')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x18, 0x18, 0x1F)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Cross-Functional Assessment | March 9, 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Internal \u2014 Confidential')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
run.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Teams: Strategy | Product | Engineering | Design | Marketing | Support & Testing')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x6E)

doc.add_page_break()

# ==========================================
# TABLE OF CONTENTS (placeholder)
# ==========================================
add_heading_styled('Table of Contents', level=1)
toc_items = [
    'Section 1: Executive Briefing',
    'Section 2: Strategy Team Report',
    'Section 3: Product Team Report',
    'Section 4: Engineering Team Report',
    'Section 5: Design Team Report',
    'Section 6: Marketing Team Report',
    'Section 7: Support & Testing Team Report',
    'Section 8: Unified 3/6/12 Month Roadmap',
    'Section 9: Team Meeting Memory & Action Items',
]
for i, item in enumerate(toc_items, 1):
    p = doc.add_paragraph(f'{item}')
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ==========================================
# SECTION 1: EXECUTIVE BRIEFING
# ==========================================
add_heading_styled('Section 1: Executive Briefing', level=1)

add_heading_styled('Meeting Purpose', level=2)
doc.add_paragraph(
    'First cross-functional team review of the Orbit platform \u2014 an AI-native autonomous fundraising '
    'intelligence platform for university advancement offices competing with GiveCampus, Givzy, EverTrue, '
    'and Blackbaud. Six teams reviewed all project deliverables: 22 files including 10 HTML prototypes '
    '(36K+ lines), a Node.js backend (130+ files), 7 Word documents, CLAUDE.md constitution, and supporting '
    'materials. The domain donororbit.com has been acquired.'
)

add_heading_styled('Overall Assessment', level=2)
add_table_from_data(
    ['Metric', 'Value'],
    [
        ['Platform Grade', 'B+ (strong foundation, significant integration gap)'],
        ['Deployment Readiness', '4.55 / 10'],
        ['Estimated Time to MVP', '8-12 weeks with focused engineering'],
        ['Backend Code Volume', '400K+ characters of service logic'],
        ['Frontend Prototype Volume', '36K+ lines across 10 HTML files'],
        ['Infrastructure Cost', '~$246/month'],
        ['Production Customers', '0 (zero)'],
        ['Recommended Pricing', '$499 - $1,999/month + Enterprise'],
    ]
)

add_heading_styled('Key Findings', level=2)
findings = [
    ('Backend is substantial: ', '400K+ characters of service logic across 20 route files, 19 service files, Docker, CI/CD, and PostgreSQL with triggers.'),
    ('Frontend is polished but disconnected: ', '36K+ lines of HTML/React prototypes using hardcoded demo data. Never connected to backend APIs.'),
    ('The Integration Chasm: ', 'Frontend and backend were built in parallel and have NEVER been connected. This is the #1 gap.'),
    ('AI agents are the differentiator: ', 'Four agents (VEO, VSO, VPGO, VCO) with Claude-powered reasoning loops. No competitor offers this.'),
    ('CRM integrations built: ', 'Salesforce NPSP, HubSpot, RE NXT with test data and test runners.'),
    ('Critical bugs exist: ', '4 runtime-breaking bugs prevent the backend from booting cleanly (auth, middleware, workers, crons).'),
    ('Test suites non-functional: ', '101+ tests exist but none can run due to schema mismatches and import errors.'),
    ('No SOC 2: ', 'Blocks university procurement. Must pursue Type I immediately.'),
    ('Zero customers: ', 'Securing 3-5 design partners via a Founders Circle program is the #1 priority.'),
    ('Strong market tailwinds: ', '$68-84T Great Wealth Transfer, 25-30% advancement staffing turnover, AI adoption acceleration.'),
]
for bold, text in findings:
    add_numbered(text, bold)

add_heading_styled('Critical Decisions Required', level=2)
decisions = [
    ('Connect vs. Rewrite frontend: ', 'Connect existing prototypes to APIs (4-6 weeks) vs. full React rewrite (3-4 months). RECOMMENDATION: Connect now, rewrite later.'),
    ('Kill Membership module: ', 'It is a different product for a different market with a different design language. Defer or spin out.'),
    ('Route all AI through backend: ', 'Frontend currently calls Claude directly (exposes API keys). One-day refactor with massive security payoff.'),
    ('Pick one launch CRM: ', 'Salesforce NPSP. Make it bulletproof. Add HubSpot in v1.1, RE NXT in v1.2.'),
    ('Build the one-minute demo path: ', 'Login \u2192 see donor \u2192 run agent \u2192 approve email \u2192 send. Every sprint judged by this.'),
]
for bold, text in decisions:
    add_numbered(text, bold)

doc.add_page_break()

# ==========================================
# SECTION 2: STRATEGY
# ==========================================
add_heading_styled('Section 2: Strategy Team Report', level=1)

add_heading_styled('Executive Summary', level=2)
doc.add_paragraph(
    'Orbit enters a market dominated by entrenched incumbents (Blackbaud at $800M+ revenue) and well-funded '
    'insurgents (EverTrue $100M+ raised, GiveCampus $170M+ raised) with zero production customers. However, '
    'no competitor offers autonomous AI agent-driven donor relationship management. This is a categorical '
    'difference, not a feature improvement.'
)

add_heading_styled('Competitive Landscape', level=2)
add_table_from_data(
    ['Competitor', 'Target Market', 'AI Capabilities', 'Key Differentiator', 'Funding/Maturity'],
    [
        ['GiveCampus', 'Higher ed (800+ institutions)', 'Limited (smart segmentation)', 'Best-in-class giving days', '$170M+ raised'],
        ['EverTrue', 'Higher ed (700+ institutions)', 'Moderate (prospect scoring, signals)', 'Data enrichment at scale', '$100M+ raised'],
        ['Gravyty/Raise', 'Higher ed + nonprofits', 'Moderate (AI email drafting)', 'First-mover in AI comms', 'Acquired by Inasmuch'],
        ['Blackbaud', 'All (dominant incumbent)', 'Minimal (basic analytics)', 'Market dominance, data lock-in', '$800M+ revenue'],
        ['Givzy', 'Nonprofits, younger donors', 'None', 'Mobile-native social giving', 'Seed stage'],
        ['ORBIT', 'Higher ed advancement', 'DEEP (4 autonomous agents)', 'AI agents manage relationships 24/7', 'Pre-revenue'],
    ]
)

add_heading_styled('Competitive Advantages', level=2)
advantages = [
    'AI-native architecture \u2014 agents ARE the product, not a feature bolted onto a CRM',
    'Full-lifecycle donor coverage with specialized agents (VEO, VSO, VPGO, VCO)',
    'Multi-CRM integration as a first-class feature (not an afterthought)',
    'Aggressive pricing: $499-$1,299/mo vs $15K-$100K/yr for incumbents',
    'Transparent AI with compliance-first guardrails (FERPA, AI disclosure, human-in-the-loop)',
    'Development velocity: v1.0 to v1.4 shipped in 5 consecutive days',
]
for a in advantages:
    add_bullet(a)

add_heading_styled('Competitive Vulnerabilities', level=2)
vulns = [
    'Zero production customers \u2014 no social proof whatsoever',
    'No SOC 2 certification \u2014 blocks most university procurement processes',
    'Frontend is prototype, not production application',
    'Single-point-of-failure on development capacity (small team)',
    '100% Anthropic API dependency (~$5,400/mo at scale) with no fallback',
    'No wealth screening integration (DonorSearch, iWave) \u2014 table stakes for advancement',
    'Reputational risk of autonomous donor contact \u2014 one bad AI email to a major donor is catastrophic',
]
for v in vulns:
    add_bullet(v)

add_heading_styled('Market Opportunity', level=2)
add_table_from_data(
    ['Metric', 'Estimate', 'Basis'],
    [
        ['TAM', '$2-4B globally', 'Higher ed advancement technology market'],
        ['SAM', '$60M-$160M annually', '3,000-4,000 institutions with budget for AI tools'],
        ['SOM (Year 3)', '$4M-$8M ARR', '100-200 institutions at $20K-$40K avg contract'],
    ]
)
doc.add_paragraph()
doc.add_paragraph(
    'Key tailwinds: Great Wealth Transfer ($68-84T intergenerational), advancement staffing crisis '
    '(25-30% annual gift officer turnover), AI adoption acceleration post-ChatGPT, stubbornly low '
    'donor retention (~19-20% for first-time donors).'
)

add_heading_styled('Strategic Positioning', level=2)
p = doc.add_paragraph()
run = p.add_run('Recommended position: ')
run.bold = True
p.add_run('"The AI Advancement Layer" \u2014 not a CRM, not a point tool. '
          '"The autonomous AI layer that sits on top of your existing CRM and turns your '
          '150-donor-per-officer ceiling into 1,500+."')

p = doc.add_paragraph()
run = p.add_run('Category creation: ')
run.bold = True
p.add_run('"Autonomous Advancement Intelligence"')

add_heading_styled('Go-To-Market Priorities (6 Months)', level=2)
add_numbered('Secure 3-5 design partner institutions (Months 1-3): Founders Circle program with free/discounted pilots in exchange for case studies and reference calls. Target mid-size institutions on Salesforce NPSP.', '1. ')
add_numbered('Achieve production frontend + SOC 2 Type I (Months 2-6): Hire senior React developer. Use Vanta/Drata to accelerate SOC 2. Ship production dashboard by Month 4.', '2. ')
add_numbered('Build CASE/AFP conference pipeline (Months 3-6): Submit speaking proposals on "AI in Advancement." Publish case studies. Produce "State of AI in Advancement" benchmark report.', '3. ')

add_heading_styled('Risk Register', level=2)
add_table_from_data(
    ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
    [
        ['No production customers within 6 months', 'HIGH', 'CRITICAL', 'Founders Circle free pilots; target small colleges first'],
        ['Anthropic API dependency (price/outage)', 'MEDIUM', 'HIGH', 'Abstraction layer for multi-LLM; expand caching beyond 40%'],
        ['Incumbent AI response (12-18 months)', 'HIGH', 'HIGH', 'Move fast on category creation; build switching costs via data'],
        ['AI outreach causes reputational damage', 'MEDIUM', 'CRITICAL', 'Human-in-the-loop enforced; shadow mode for first 30 days'],
        ['Team capacity cannot scale', 'HIGH', 'HIGH', 'Hire React dev + CS person; raise pre-seed $500K-$1.5M'],
    ]
)

doc.add_page_break()

# ==========================================
# SECTION 3: PRODUCT
# ==========================================
add_heading_styled('Section 3: Product Team Report', level=1)

add_heading_styled('Executive Summary', level=2)
doc.add_paragraph(
    'The backend is more real than it appears \u2014 20 route files, 19 service files, Docker, CI/CD, '
    'PostgreSQL with triggers and constraints. The frontend exists as 36K+ lines of polished HTML prototypes '
    'using CDN React + Babel. The core gap is an "integration chasm" where frontend and backend were built '
    'in parallel without being wired together. The founder is right: a lot was built. The issue is not volume '
    'but connectivity.'
)

add_heading_styled('Module Maturity Assessment', level=2)
add_table_from_data(
    ['Module', 'UI Ready (1-5)', 'Backend Ready (1-5)', 'Integration (1-5)', 'Notes'],
    [
        ['Dashboard', '4', '3', '1', 'Beautiful UI, hardcoded data'],
        ['Donors', '3', '3', '1', 'Both sides exist, not connected'],
        ['Gifts', '3', '4', '1', 'Backend strong (triggers, matching)'],
        ['Agents', '4', '4', '2', 'Most mature module overall'],
        ['Campaigns', '3', '2', '1', 'Basic CRUD both sides'],
        ['Outreach', '3', '3', '1', 'Compose + delivery exist'],
        ['Analytics', '2', '3', '1', 'Backend has engines, no UI routes'],
        ['Integrations', '4', '4', '2', 'Closest to end-to-end'],
        ['Membership', '3', '1', '0', 'DIFFERENT PRODUCT \u2014 defer'],
        ['Comms Hub', '4', '1', '0', 'Two polished builders, zero backend'],
        ['Superadmin', '3', '4', '2', 'Thorough backend, org provisioning'],
        ['Auth', '2', '4', '3', 'Backend solid, no token management UI'],
        ['Billing', '1', '4', '2', 'Stripe lifecycle built, no UI'],
    ]
)

add_heading_styled('Feature-Functionality Gap Analysis', level=2)
p = doc.add_paragraph()
run = p.add_run('Frontend-only features (UI built, no backend): ')
run.bold = True
p.add_run('Membership module, Communications Hub email builders, Day of Giving gamification, '
          'Agent personality/memory system, 30+ email/SMS provider directory')

p = doc.add_paragraph()
run = p.add_run('Backend-only capabilities (no UI): ')
run.bold = True
p.add_run('Donor Intelligence (37K chars), Predictive Engine (26K chars), VSO Engine (48K chars \u2014 largest file), '
          'Planned Giving Brain (32K chars), Signal Ingestion (16K chars), Marketing Studio (18K chars), '
          'Gift Write-Back (22K chars), Billing/Subscription system')

add_heading_styled('Critical Path to MVP (8-12 Weeks)', level=2)
mvp_items = [
    'Frontend-backend API wiring \u2014 3-4 weeks (every component reads hardcoded data today)',
    'Auth flow completion \u2014 1 week (backend has JWT, frontend has login form, need: token storage, refresh, protected routes)',
    'CRM initial sync trigger \u2014 1-2 weeks (trigger first sync from integration wizard, show progress)',
    'Outreach send pipeline \u2014 1-2 weeks (connect agent decision to actual email/SMS delivery)',
    'Settings for API keys \u2014 customers need to enter their own Anthropic/SendGrid/Twilio keys',
]
for item in mvp_items:
    add_bullet(item)

add_heading_styled('Founder\'s Question Answered', level=2)
p = doc.add_paragraph()
run = p.add_run('"I feel we built a lot and not all is functional or used, but may be wrong."')
run.italic = True
doc.add_paragraph(
    'You are right on both counts. You built far more than most startups at this stage. Features (things users '
    'can see) are abundant and polished. Functionality (things that work end-to-end) is near zero. The bridge '
    'is integration work \u2014 plumbing, not glamour. For every feature in the UI, you need: (1) an API call '
    'that fetches real data, (2) error handling, (3) loading states, (4) mutation endpoints. Your backend '
    'services are rich \u2014 donorIntelligence.js (37K chars), predictiveEngine.js (26K chars), vsoEngine.js '
    '(48K chars) represent real intellectual property. But they are underwater \u2014 no user can see or benefit '
    'from them. RECOMMENDATION: Wire first, build later. You have more than enough features for MVP.'
)

doc.add_page_break()

# ==========================================
# SECTION 4: ENGINEERING
# ==========================================
add_heading_styled('Section 4: Engineering Team Report', level=1)

add_heading_styled('Executive Summary', level=2)
doc.add_paragraph(
    '~60% of backend code is functional and well-structured. ~40% consists of services referencing non-existent '
    'modules/tables, incompatible DB access patterns, and non-functional test suites. The project needs 4-8 weeks '
    'of focused engineering to reach production stability.'
)

add_heading_styled('Critical Bugs Found', level=2)
add_table_from_data(
    ['Bug', 'File', 'Impact', 'Fix Effort'],
    [
        ['loginLimiter used before definition', 'src/routes/auth.js', 'Login route crashes at import', '1 hour'],
        ['Middleware org_id mismatch (orgId vs org_id)', 'auth.js + tenant.js', 'Tenant isolation broken', '2 hours'],
        ['Nested cron re-registration', 'src/jobs/index.js', 'N duplicate jobs after N days', '2 hours'],
        ['Workers use Knex, DB uses raw pg Pool', 'src/workers/index.ts', 'Entire async pipeline broken', '1-2 days'],
    ]
)

add_heading_styled('Security Review', level=2)
add_table_from_data(
    ['Area', 'Rating (1-5)', 'Notes'],
    [
        ['Authentication', '3', 'JWT + refresh rotation good. loginLimiter broken. dev-secret fallback dangerous.'],
        ['Authorization', '4', 'RBAC hierarchy, IDOR prevention, impersonation tokens (1hr TTL)'],
        ['Input Validation', '3', 'AI jailbreak detection good. No schema validation library (Zod/Joi).'],
        ['SQL Injection', '5', 'All queries parameterized. No string concatenation in SQL.'],
        ['Data Encryption', '2', 'Integration creds encrypted. No key rotation. Hardcoded seed hash.'],
        ['Rate Limiting', '3', 'Global + AI limits work. Login limiter broken. No limits on billing/gifts.'],
        ['Secrets Management', '2', 'Env vars only. 30+ secrets. No vault or rotation strategy.'],
    ]
)

add_heading_styled('API Completeness', level=2)
doc.add_paragraph(
    '8 route files present and mostly functional. 8+ additional route files referenced in server.js but missing '
    '(tenant, vso, giving, payment, plannedGiving, webhooks, outreach, campaigns). The API is approximately 50% '
    'complete by endpoint count.'
)

add_heading_styled('Test Coverage', level=2)
doc.add_paragraph(
    'Effective test coverage: ~0%. 101+ tests exist across 10 files. None can run successfully due to import '
    'mismatches (tests import function names that helpers doesn\'t export) and schema inconsistencies (tests '
    'reference table "orgs" but schema defines "organizations"; tests reference columns first_name/last_name '
    'but schema has "name"). Tests appear written against the CLAUDE.md specification rather than the actual '
    'implementation.'
)

add_heading_styled('Deployment Readiness: 4.55/10', level=2)
doc.add_paragraph(
    'Docker and CI/CD infrastructure is excellent (5-job GitHub Actions with Trivy scanning, staged deploys). '
    'Application code has 4 runtime-breaking bugs that prevent basic smoke testing. The pipeline would fail '
    'at the test stage, blocking all downstream jobs.'
)

add_heading_styled('Top 10 Engineering Priorities', level=2)
eng_priorities = [
    'Fix auth.js loginLimiter bug (1 hour)',
    'Resolve middleware org_id naming mismatch (2 hours)',
    'Eliminate dual entry points \u2014 delete root index.js and server.js (2 hours)',
    'Fix workers DB layer \u2014 rewrite to use raw pg Pool (1-2 days)',
    'Fix test suites \u2014 align with actual schema (2-3 days)',
    'Reconcile table/column names across codebase (1 day)',
    'Fix nested cron re-registration in jobs (2 hours)',
    'Consolidate AI client implementations \u2014 standardize on Anthropic SDK (4 hours)',
    'Add input validation library (Zod) (2 days)',
    'Create missing route files or remove dead mounts from server.js (variable)',
]
for i, p_text in enumerate(eng_priorities, 1):
    add_numbered(f'{p_text}', f'{i}. ')

doc.add_page_break()

# ==========================================
# SECTION 5: DESIGN
# ==========================================
add_heading_styled('Section 5: Design Team Report', level=1)

add_heading_styled('Executive Summary', level=2)
doc.add_paragraph(
    'The prototype suite is remarkably ambitious \u2014 10 HTML files totaling 30K+ lines. The dashboard\'s '
    'design token foundation is solid with a warm, professional aesthetic (teal accents, Plus Jakarta Sans, '
    'Outfit headings). However, significant inconsistencies exist across modules, and the Membership module '
    'uses an entirely different visual language (dark navy/gold). Responsive design is essentially absent. '
    'Accessibility has not been addressed \u2014 a potential deal-breaker for university procurement (ADA/Section 508).'
)

add_heading_styled('Design System Audit', level=2)
add_table_from_data(
    ['Token', 'Dashboard (canonical)', 'Agents', 'Membership', 'Issue'],
    [
        ['--bg', '#f6f5f1 (warm)', '#f0f2f5 (cool)', '#0d1b2e (navy)', 'Membership is completely different product'],
        ['--teal', '#2a8c7e', '#2a8c7e', '#c9973a (gold)', 'Membership uses gold instead of teal'],
        ['--ink', '#18181f', '#18181f', 'white', 'Inverted color scheme in Membership'],
        ['Body font', 'Plus Jakarta Sans', 'Plus Jakarta Sans', 'DM Sans', 'Different type family'],
        ['Display font', 'Outfit', 'Outfit', 'Playfair Display', 'Serif vs sans-serif \u2014 different personality'],
    ]
)

add_heading_styled('UX Assessment by Module', level=2)
add_table_from_data(
    ['Module', 'Visual Polish', 'Usability', 'Info Density', 'Accessibility', 'Notes'],
    [
        ['Dashboard', '4/5', '3/5', '4/5', '2/5', 'Strongest module. Icon-rail sidebar is smart.'],
        ['Agent Center', '4/5', '4/5', '3/5', '2/5', 'ChatGPT-style chat is familiar and effective.'],
        ['Comms v2', '4/5', '4/5', '3/5', '2/5', 'Excellent token system with 40+ merge tokens.'],
        ['Giving Portal', '3/5', '3/5', '3/5', '2/5', 'Clean donor-facing page. Has mobile meta tags.'],
        ['Membership', '3/5', '3/5', '3/5', '1/5', 'Dark theme is visually striking but jarring in context.'],
        ['Superadmin', '3/5', '4/5', '4/5', '2/5', 'Onboarding wizard with AI chat is genuinely clever.'],
    ]
)

add_heading_styled('Persona-UI Fit', level=2)
persona_fit = [
    ('Sarah (Major Gift Officer): ', 'Strong fit. Needs a "My Portfolio" aggregate view of her prospects, tasks, signals, and recommendations.'),
    ('Marcus (Annual Giving Director): ', 'Good fit. Needs campaign-level analytics and segment performance breakdowns.'),
    ('Dani (CRM Admin): ', 'Good fit. Needs a "System Health" dashboard showing sync status, error logs, API rate limits.'),
    ('Priya (Stewardship Officer): ', 'WEAKEST FIT. No dedicated stewardship module exists. Critical gap for donor retention.'),
]
for bold, text in persona_fit:
    add_bullet(text, bold)

add_heading_styled('Top 10 Design Priorities', level=2)
design_priorities = [
    'Unify design token system into single shared CSS file',
    'Standardize sidebar navigation \u2014 adopt collapsible icon-rail everywhere',
    'Resolve Membership module visual identity (bring into core theme or spin out)',
    'Build shared React component library (Button, Card, Modal, Table, Form, Badge)',
    'Add stewardship module for Priya persona',
    'Implement core accessibility (ARIA landmarks, focus indicators, contrast fixes)',
    'Create "My Portfolio" view for gift officers',
    'Simplify agent config with presets (Professional, Warm, Direct)',
    'Add campaign-level analytics for Annual Giving Directors',
    'Responsive design for key mobile surfaces (giving portal, agent chat, approval queue)',
]
for i, p_text in enumerate(design_priorities, 1):
    add_numbered(f'{p_text}', f'{i}. ')

doc.add_page_break()

# ==========================================
# SECTION 6: MARKETING
# ==========================================
add_heading_styled('Section 6: Marketing Team Report', level=1)

add_heading_styled('Executive Summary', level=2)
doc.add_paragraph(
    'Marketing readiness is at ~40-50%. The bones are there \u2014 clean brand name, domain secured (donororbit.com), '
    'landing page and pitch deck created. Critical gaps in social proof, case studies, SOC 2, and conference '
    'presence must be closed before going to market.'
)

add_heading_styled('Brand Assessment', level=2)
doc.add_paragraph(
    '"Orbit" / "Donor Orbit" \u2014 short, memorable, and evocative. Implies gravitational pull (keeping donors '
    'in your orbit), continuous motion, and systematic ongoing relationships. donororbit.com is solid. '
    'RECOMMENDATION: File trademark applications for "Orbit" and "Donor Orbit" in SaaS/fundraising classes '
    '(Class 42/36) immediately. Secure defensive domains (orbithq.com, getorbit.com, orbitfundraising.com).'
)

add_heading_styled('Positioning Statement', level=2)
p = doc.add_paragraph()
run = p.add_run(
    'For small-to-mid-size college and university advancement offices that are understaffed and under-resourced, '
    'Orbit is the AI-native fundraising platform that autonomously researches prospects, prioritizes outreach, '
    'and generates personalized donor communications \u2014 so your team can raise more money with fewer people.'
)
run.italic = True

add_heading_styled('Pricing Recommendation', level=2)
add_table_from_data(
    ['Tier', 'Price', 'Target', 'Includes'],
    [
        ['Starter', '$499/mo', 'Community colleges, 1-3 staff', 'Core AI, 10K records, 1 CRM, email support'],
        ['Growth', '$999/mo', 'Small colleges, 3-8 staff', 'Full AI suite, 50K records, 2 CRMs, priority support'],
        ['Pro', '$1,999/mo', 'Mid-size, 8-15 staff', 'Unlimited records, full integrations, dedicated CSM'],
        ['Enterprise', 'Custom', 'Large institutions, systems', 'Multi-campus, SLA, advanced security, QBRs'],
    ]
)

add_heading_styled('Competitive Messaging', level=2)
messaging = [
    ('vs. GiveCampus: ', '"GiveCampus is great for giving days. Orbit is what happens the other 364 days of the year."'),
    ('vs. Blackbaud: ', '"You spent $50K on Raiser\'s Edge. Orbit is the AI layer that finally gets your team to use it."'),
    ('vs. EverTrue: ', '"EverTrue shows you data. Orbit acts on it."'),
    ('vs. Givzy: ', '"Givzy makes it easy to give. Orbit makes it easy to fundraise."'),
]
for bold, text in messaging:
    add_bullet(text, bold)

add_heading_styled('Channel Strategy', level=2)
add_table_from_data(
    ['Channel', 'Priority', 'Approach'],
    [
        ['CASE Conferences', 'CRITICAL', 'Sponsor, exhibit, present sessions. Annual + District + DRIVE.'],
        ['LinkedIn', 'HIGH', 'Thought leadership posts, targeted ads by job title + higher ed'],
        ['Direct Outreach', 'CRITICAL', 'Cold email + LinkedIn to advancement leaders. List is finite (~4,000 institutions).'],
        ['Consultant Channel', 'HIGH', 'Build relationships with Grenzebach Glier, CCS, BWF, Marts & Lundy'],
        ['AFP ICON', 'HIGH', 'Broader nonprofit audience but includes higher ed'],
        ['Podcasts', 'MEDIUM', 'Guest appearances on advancement-focused shows'],
        ['Peer Recommendations', 'CRITICAL', 'Every pilot customer must become a reference'],
    ]
)

doc.add_page_break()

# ==========================================
# SECTION 7: SUPPORT & TESTING
# ==========================================
add_heading_styled('Section 7: Support & Testing Team Report', level=1)

add_heading_styled('Executive Summary', level=2)
doc.add_paragraph(
    'Testing infrastructure exists in form but not function. 101+ tests across 10 files cannot execute due to '
    'import mismatches and schema inconsistencies. CI/CD pipeline design is excellent but would fail at the test '
    'stage. Deployment is blocked by 12 specific launch blockers.'
)

add_heading_styled('Launch Blockers', level=2)
blockers = [
    'Fix auth.js loginLimiter runtime crash',
    'Fix middleware org_id naming mismatch (breaks tenant isolation)',
    'Fix or rewrite test suites to match actual schema',
    'Add health check endpoint verification',
    'SOC 2 Type I (minimum for university procurement)',
    'FERPA compliance documentation',
    'Data backup and recovery procedures',
    'Monitoring and alerting setup (no Sentry/Datadog configured)',
    'Error handling standardization across routes',
    'Rate limiting on all sensitive endpoints',
    'Privacy policy and data processing agreements',
    'Customer onboarding documentation',
]
for i, b in enumerate(blockers, 1):
    add_numbered(f'{b}', f'{i}. ')

add_heading_styled('Monitoring Gaps', level=2)
doc.add_paragraph(
    'No application performance monitoring (APM), no error tracking service (Sentry), no uptime monitoring, '
    'no alerting. Winston logging exists but with no log aggregation or search. These are must-haves before '
    'any customer touches the platform.'
)

add_heading_styled('Support Infrastructure Needed', level=2)
support_items = [
    'Help documentation / knowledge base',
    'Ticketing system (Intercom, Zendesk, or similar)',
    'SLA framework and escalation procedures',
    'Customer onboarding playbook (first 30 days)',
    'Status page for uptime transparency',
]
for item in support_items:
    add_bullet(item)

doc.add_page_break()

# ==========================================
# SECTION 8: ROADMAP
# ==========================================
add_heading_styled('Section 8: Unified 3/6/12 Month Roadmap', level=1)

# Phase 1
add_heading_styled('Phase 1: Foundation & First Customers (Months 1-3)', level=2)

add_heading_styled('Month 1: "Make It Work"', level=3)
month1 = [
    ('ENGINEERING: ', 'Fix 4 critical runtime bugs (Week 1). Eliminate dual entry points. Wire frontend to backend APIs for auth, donors, and agent run (Weeks 2-4). Fix test suites and get CI green.'),
    ('PRODUCT: ', 'Define MVP scope. Cut Membership module. Route all AI through backend proxy.'),
    ('DESIGN: ', 'Unify design tokens. Standardize sidebar navigation.'),
    ('MARKETING: ', 'Launch donororbit.com. Publish 5 blog posts. LinkedIn presence. Begin outreach to 15 design partner prospects.'),
    ('SUPPORT: ', 'Set up Sentry. Draft privacy policy and terms of service. Basic help docs.'),
]
for bold, text in month1:
    add_bullet(text, bold)

add_heading_styled('Month 2: "Make It Real"', level=3)
month2 = [
    ('ENGINEERING: ', 'Complete API wiring for MVP modules. Salesforce NPSP end-to-end test. Email send via SendGrid working. Agent \u2192 email \u2192 approval \u2192 send pipeline complete.'),
    ('PRODUCT: ', 'Internal dogfooding. "One-minute demo path" working. Begin billing UI.'),
    ('DESIGN: ', 'Build 10 core components or polish API-connected prototypes. Accessibility pass.'),
    ('MARKETING: ', 'First design partner signed. First webinar. CASE speaking proposals submitted.'),
    ('SUPPORT: ', 'Monitoring and alerting live. Onboarding playbook. SOC 2 Type I started.'),
]
for bold, text in month2:
    add_bullet(text, bold)

add_heading_styled('Month 3: "First Customer Live"', level=3)
month3 = [
    ('ENGINEERING: ', 'First design partner onboarded with real CRM data. CRM sync running. Agent generating real outreach. Performance testing.'),
    ('PRODUCT: ', 'First real user feedback. Iterate on UX. Add SMS via Twilio.'),
    ('DESIGN: ', 'Iterate on feedback. Add "My Portfolio" view for gift officers.'),
    ('MARKETING: ', 'First case study draft. Attend CASE District conference. 2nd/3rd partners signed.'),
    ('SUPPORT: ', 'First support experience documented. SLA drafted. Backup/recovery tested.'),
]
for bold, text in month3:
    add_bullet(text, bold)

# Phase 2
add_heading_styled('Phase 2: Growth & Validation (Months 4-6)', level=2)

add_heading_styled('Month 4: "Prove the Value"', level=3)
month4 = [
    'HubSpot integration launch',
    'VSO agent (stewardship) activated',
    'Pledge management and campaign tracking',
    'Predictive engine exposed in dashboard',
    '3 design partners active, collecting metrics',
    '2 case studies published',
]
for item in month4:
    add_bullet(item)

add_heading_styled('Month 5: "Start Selling"', level=3)
month5 = [
    'Communications Hub backend (templates, sequences, tokens)',
    'RE NXT integration (3rd CRM)',
    'VCO agent (campaigns) activated',
    'Billing portal live with plan gating',
    '5 paying customers target',
    'Press outreach to Inside Higher Ed / Chronicle',
]
for item in month5:
    add_bullet(item)

add_heading_styled('Month 6: "Launch"', level=3)
month6 = [
    'Performance and security hardening',
    'Input validation library deployed (Zod)',
    'Full analytics dashboard with real predictive data',
    'SOC 2 Type I complete',
    '10 paying customers target',
    'CASE Annual conference presence',
    '3 published case studies with ROI metrics',
    'Formal launch announcement',
]
for item in month6:
    add_bullet(item)

# Phase 3
add_heading_styled('Phase 3: Scale & Differentiate (Months 7-12)', level=2)

add_heading_styled('Months 7-9: "Expand the Platform"', level=3)
phase3a = [
    'VPGO agent (planned giving) with compliance workflow',
    'Wealth screening integration (DonorSearch or iWave)',
    'Day of Giving gamification backend',
    'React frontend migration (if not started)',
    'Multi-campus / university system support',
    'SOC 2 Type II process initiated',
    '25 paying customers target',
    'Series Seed fundraise ($500K-$1.5M)',
]
for item in phase3a:
    add_bullet(item)

add_heading_styled('Months 10-12: "Category Leader"', level=3)
phase3b = [
    'AI model fine-tuning on institutional data',
    'Agent memory persistence (learns from officer edits)',
    'Donor self-service portal',
    'API for third-party integrations',
    'Advanced security (SSO, SAML, MFA)',
    'SOC 2 Type II complete',
    '50+ paying customers target',
    '$500K+ ARR milestone',
    'Begin hiring: CS manager, sales rep, frontend developer',
]
for item in phase3b:
    add_bullet(item)

doc.add_page_break()

# ==========================================
# SECTION 9: MEETING MEMORY
# ==========================================
add_heading_styled('Section 9: Team Meeting Memory & Action Items', level=1)

add_heading_styled('Meeting Record', level=2)
add_table_from_data(
    ['Field', 'Value'],
    [
        ['Date', 'March 9, 2026'],
        ['Type', 'First Cross-Functional Team Review'],
        ['Participants', 'Strategy, Product, Engineering, Design, Marketing, Support/Testing'],
        ['Files Reviewed', '22 deliverables (10 HTML, 7 DOCX, backend tarball, PPTX, 3 MD files)'],
        ['Platform Version', 'Backend v1.4.0 / Dashboard v4.0.0'],
    ]
)

add_heading_styled('Consensus Decisions', level=2)
consensus = [
    'Platform grade: B+ (strong foundation, integration gap)',
    'MVP timeline: 8-12 weeks with focused engineering',
    'Priority #1: Wire frontend to backend (close the integration chasm)',
    'Launch CRM: Salesforce NPSP (make it bulletproof, add others later)',
    'Membership module: Defer (different product, different market)',
    'AI routing: All Claude calls through backend proxy (security requirement)',
    'Frontend strategy: Connect existing prototypes now, React rewrite later',
    'Pricing: $499 / $999 / $1,999/mo tiers + enterprise custom',
    'First customers: 3-5 design partners via Founders Circle program',
    'Category name: "Autonomous Advancement Intelligence"',
]
for i, c in enumerate(consensus, 1):
    add_numbered(f'{c}', f'{i}. ')

add_heading_styled('Open Questions for Founder', level=2)
questions = [
    'Fundraising plans? A seed round ($500K-$1.5M) is needed for team hiring.',
    'Target first hire \u2014 senior React developer or customer success person?',
    'Is the Greenfield University pilot data simulated or from a real engagement?',
    'Budget for CASE conference presence (booth ~$5-10K, travel ~$3-5K)?',
    'SOC 2 timeline preference \u2014 Type I first (faster) or go directly to Type II?',
]
for q in questions:
    add_bullet(q)

add_heading_styled('Immediate Action Items (Week 1)', level=2)
immediate = [
    'Fix 4 critical backend bugs (auth, middleware, workers, crons)',
    'Unify design tokens into single CSS file',
    'Launch donororbit.com (set up Google Workspace email first)',
    'Begin design partner outreach \u2014 identify 15 target institutions',
    'File trademark applications for "Orbit" and "Donor Orbit"',
]
for item in immediate:
    add_bullet(item)

add_heading_styled('Key Metrics to Track', level=2)
add_table_from_data(
    ['Metric', 'Target', 'Frequency'],
    [
        ['Time to first paying customer', '< 90 days from MVP', 'Weekly'],
        ['Design partner conversion rate', '> 60% of pilots', 'Monthly'],
        ['Agent outreach approval rate', '> 85% approved as-is', 'Weekly'],
        ['Donor response rate to AI comms', '> 25% open rate', 'Weekly'],
        ['CRM sync reliability', '> 99.5% uptime', 'Daily'],
        ['Infrastructure cost per customer', '< $50/customer/mo', 'Monthly'],
        ['Claude token cost per donor/month', '< $0.15/donor', 'Monthly'],
        ['Test suite pass rate', '100% on CI', 'Per commit'],
    ]
)

# Save
output_path = '/Users/colin/Documents/Orbit_Team_Review_Report.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
